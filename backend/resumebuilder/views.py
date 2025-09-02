from rest_framework.views import APIView
from rest_framework.permissions import IsAuthenticated, AllowAny
from rest_framework.response import Response
from rest_framework import status
from django.http import HttpResponse
from .models import ResumeSnapshot
from .serializers import ResumeSnapshotSerializer
from io import BytesIO
try:
    from docx import Document
except Exception:
    Document = None
try:
    from playwright.sync_api import sync_playwright
except Exception:
    sync_playwright = None


class ResumeExportView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request):
        """Exports resume in requested format. Supports html echo, and doc/docx (text-based)."""
        export_format = (request.query_params.get('format') or request.data.get('format') or 'pdf').lower()
        html = request.data.get('html')
        data = request.data.get('data') or {}
        # HTML echo for quick preview
        if export_format == 'html':
            if not html:
                return Response({'detail': 'html is required'}, status=400)
            return HttpResponse(html, content_type='text/html')
        # Text-based DOCX/DOC from resume data
        if export_format in ('docx', 'doc'):
            if Document is None:
                return Response({'detail': 'python-docx not installed on server'}, status=501)
            resume = data.get('resumeData') or {}
            doc = Document()
            p = doc.add_paragraph()
            run = p.add_run(resume.get('personal', {}).get('fullName', '') or 'Resume')
            run.bold = True
            # Helper to add section
            def add_heading(text):
                doc.add_paragraph().add_run('')
                h = doc.add_paragraph().add_run(text)
                h.bold = True
            def add_text(text):
                doc.add_paragraph(str(text or ''))
            personal = resume.get('personal', {})
            add_text(personal.get('title'))
            add_text(' | '.join([x for x in [personal.get('email'), personal.get('phone'), personal.get('location')] if x]))
            if personal.get('summary'):
                add_heading('Summary')
                add_text(personal.get('summary'))
            exps = resume.get('experience') or []
            if exps:
                add_heading('Experience')
                for e in exps:
                    add_text(f"{e.get('title','')} — {e.get('company','')}")
                    add_text(f"{e.get('startDate','')} - {('Present' if e.get('current') else e.get('endDate',''))}")
                    if e.get('description'):
                        add_text(e.get('description'))
            edus = resume.get('education') or []
            if edus:
                add_heading('Education')
                for ed in edus:
                    add_text(f"{ed.get('degree','')}{(' in ' + ed.get('field','')) if ed.get('field') else ''}")
                    add_text(f"{ed.get('institution','')}{(' • ' + ed.get('year','')) if ed.get('year') else ''}{(' • GPA: ' + ed.get('gpa','')) if ed.get('gpa') else ''}")
            skills = resume.get('skills') or []
            if skills:
                add_heading('Skills')
                add_text(', '.join(skills))
            bio = BytesIO()
            doc.save(bio)
            bio.seek(0)
            content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            filename = 'resume.docx'
            if export_format == 'doc':
                # Old .doc content-type; still a docx file but hints at .doc
                content_type = 'application/msword'
                filename = 'resume.doc'
            resp = HttpResponse(bio.getvalue(), content_type=content_type)
            resp['Content-Disposition'] = f'attachment; filename="{filename}"'
            return resp
        # Server-side PDF (selectable text) via Playwright, when available
        if export_format == 'pdf':
            html_doc = request.data.get('html')
            if not html_doc:
                return Response({'detail': 'html is required for server-side PDF export'}, status=400)
            if sync_playwright is None:
                return Response({'detail': 'Server-side PDF not available (Playwright not installed)'}, status=501)
            try:
                with sync_playwright() as p:
                    browser = p.chromium.launch()
                    context = browser.new_context()
                    page = context.new_page()
                    # Set default viewport to A4 at 96 DPI approx (794x1123 px)
                    page.set_content(html_doc, wait_until="networkidle")
                    pdf_bytes = page.pdf(format="A4", print_background=True, margin={"top": "0", "right": "0", "bottom": "0", "left": "0"})
                    context.close()
                    browser.close()
                resp = HttpResponse(pdf_bytes, content_type='application/pdf')
                resp['Content-Disposition'] = 'attachment; filename="resume.pdf"'
                return resp
            except Exception as e:
                return Response({'detail': f'PDF generation failed: {str(e)}'}, status=500)
        # Unknown/unsupported format
        return Response({'detail': f'Unsupported export format: {export_format}'}, status=400)


class ResumeTemplatesView(APIView):
    permission_classes = [AllowAny]

    def get(self, request):
        # Expose available template ids used by frontend
        # Only include templates that are actually implemented in LivePreview.jsx
        return Response({
            'templates': [
                {'id': 'modern', 'name': 'Modern Professional'},
                {'id': 'classic', 'name': 'Classic Executive'},
                {'id': 'minimal', 'name': 'Minimal Clean'},
                {'id': 'technical', 'name': 'Technical Pro'},
                {'id': 'elegant', 'name': 'Elegant Serif'},
                {'id': 'compact', 'name': 'Compact One-Page'},
            ]
        })


class ResumeSnapshotView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request):
        qs = ResumeSnapshot.objects.filter(user=request.user)
        return Response({'items': ResumeSnapshotSerializer(qs, many=True).data})

    def post(self, request):
        title = (request.data.get('title') or 'Resume').strip() or 'Resume'
        data = request.data.get('data') or {}
        snap = ResumeSnapshot.objects.create(user=request.user, title=title, data=data)
        return Response(ResumeSnapshotSerializer(snap).data, status=201)


class ResumeSnapshotDetailView(APIView):
    permission_classes = [IsAuthenticated]

    def get_object(self, request, pk):
        try:
            return ResumeSnapshot.objects.get(pk=pk, user=request.user)
        except ResumeSnapshot.DoesNotExist:
            return None

    def patch(self, request, pk):
        snap = self.get_object(request, pk)
        if not snap:
            return Response({'detail': 'Not found'}, status=404)
        title = request.data.get('title')
        data = request.data.get('data')
        if title is not None:
            title = (title or '').strip() or snap.title
            snap.title = title
        if data is not None:
            snap.data = data
        snap.save()
        return Response(ResumeSnapshotSerializer(snap).data)

    def delete(self, request, pk):
        snap = self.get_object(request, pk)
        if not snap:
            return Response({'detail': 'Not found'}, status=404)
        snap.delete()
        return Response(status=204)
